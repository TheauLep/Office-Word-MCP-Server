"""
Document utility functions for Word Document Server.
"""
import json
from typing import Dict, List, Any
from docx import Document


def get_document_properties(doc_path: str) -> Dict[str, Any]:
    """Get properties of a Word document."""
    import os
    if not os.path.exists(doc_path):
        return {"error": f"Document {doc_path} does not exist"}
    
    try:
        doc = Document(doc_path)
        core_props = doc.core_properties
        
        return {
            "title": core_props.title or "",
            "author": core_props.author or "",
            "subject": core_props.subject or "",
            "keywords": core_props.keywords or "",
            "created": str(core_props.created) if core_props.created else "",
            "modified": str(core_props.modified) if core_props.modified else "",
            "last_modified_by": core_props.last_modified_by or "",
            "revision": core_props.revision or 0,
            "page_count": len(doc.sections),
            "word_count": sum(len(paragraph.text.split()) for paragraph in doc.paragraphs),
            "paragraph_count": len(doc.paragraphs),
            "table_count": len(doc.tables)
        }
    except Exception as e:
        return {"error": f"Failed to get document properties: {str(e)}"}


def extract_document_text(doc_path: str) -> str:
    """Extract all text from a Word document."""
    import os
    if not os.path.exists(doc_path):
        return f"Document {doc_path} does not exist"
    
    try:
        doc = Document(doc_path)
        text = []
        
        for paragraph in doc.paragraphs:
            text.append(paragraph.text)
            
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        text.append(paragraph.text)
        
        return "\n".join(text)
    except Exception as e:
        return f"Failed to extract text: {str(e)}"


def get_document_structure(doc_path: str) -> Dict[str, Any]:
    """Get the structure of a Word document."""
    import os
    if not os.path.exists(doc_path):
        return {"error": f"Document {doc_path} does not exist"}
    
    try:
        doc = Document(doc_path)
        structure = {
            "paragraphs": [],
            "tables": []
        }
        
        # Get paragraphs
        for i, para in enumerate(doc.paragraphs):
            structure["paragraphs"].append({
                "index": i,
                "text": para.text[:100] + ("..." if len(para.text) > 100 else ""),
                "style": para.style.name if para.style else "Normal"
            })
        
        # Get tables
        for i, table in enumerate(doc.tables):
            table_data = {
                "index": i,
                "rows": len(table.rows),
                "columns": len(table.columns),
                "preview": []
            }
            
            # Get sample of table data
            max_rows = min(3, len(table.rows))
            for row_idx in range(max_rows):
                row_data = []
                max_cols = min(3, len(table.columns))
                for col_idx in range(max_cols):
                    try:
                        cell_text = table.cell(row_idx, col_idx).text
                        row_data.append(cell_text[:20] + ("..." if len(cell_text) > 20 else ""))
                    except IndexError:
                        row_data.append("N/A")
                table_data["preview"].append(row_data)
            
            structure["tables"].append(table_data)
        
        return structure
    except Exception as e:
        return {"error": f"Failed to get document structure: {str(e)}"}


def find_paragraph_by_text(doc, text, partial_match=False):
    """
    Find paragraphs containing specific text.
    
    Args:
        doc: Document object
        text: Text to search for
        partial_match: If True, matches paragraphs containing the text; if False, matches exact text
        
    Returns:
        List of paragraph indices that match the criteria
    """
    matching_paragraphs = []
    
    for i, para in enumerate(doc.paragraphs):
        if partial_match and text in para.text:
            matching_paragraphs.append(i)
        elif not partial_match and para.text == text:
            matching_paragraphs.append(i)
            
    return matching_paragraphs


def find_and_replace_text(doc, old_text, new_text):
    """
    Find and replace text throughout the document.
    
    This function properly handles text that spans across multiple runs (formatted text segments).
    
    Args:
        doc: Document object
        old_text: Text to find
        new_text: Text to replace with
        
    Returns:
        Number of replacements made
    """
    count = 0
    
    # Helper function to replace text in a paragraph that may span multiple runs
    def replace_in_paragraph(para):
        replacements = 0
        paragraph_text = para.text
        
        if old_text not in paragraph_text:
            return 0
        
        # Count how many replacements we'll make
        replacement_count = paragraph_text.count(old_text)
        if replacement_count == 0:
            return 0
        
        # If the text is contained within a single run, use simple replacement
        simple_replacement_done = False
        for run in para.runs:
            if old_text in run.text:
                # Count occurrences in this run
                run_count = run.text.count(old_text)
                run.text = run.text.replace(old_text, new_text)
                replacements += run_count
                simple_replacement_done = True
        
        # If simple replacement worked (text didn't span runs), return
        if simple_replacement_done and old_text not in para.text:
            return replacements
        
        # Handle complex case where text spans multiple runs
        # Use a simpler but more reliable approach for cross-run replacement
        full_text = para.text
        
        # If we still have text to replace (meaning it spans runs), use paragraph-level replacement
        if old_text in full_text:
            new_full_text = full_text.replace(old_text, new_text)
            replacement_count = full_text.count(old_text)
            
            # Clear all runs and create a single run with the new text
            # This preserves the replacement but loses some formatting
            for run in para.runs:
                run.clear()
            
            # Remove all runs except the first one
            while len(para.runs) > 1:
                para._element.remove(para.runs[-1]._element)
            
            # Set the text in the first run
            if para.runs:
                para.runs[0].text = new_full_text
            else:
                para.add_run(new_full_text)
            
            replacements += replacement_count
        
        return replacements
    
    # Search in paragraphs
    for para in doc.paragraphs:
        count += replace_in_paragraph(para)
    
    # Search in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    count += replace_in_paragraph(para)
    
    return count


def get_document_xml(doc_path: str) -> str:
    """Extract and return the raw XML structure of the Word document (word/document.xml)."""
    import os
    import zipfile
    if not os.path.exists(doc_path):
        return f"Document {doc_path} does not exist"
    try:
        with zipfile.ZipFile(doc_path) as docx_zip:
            with docx_zip.open('word/document.xml') as xml_file:
                return xml_file.read().decode('utf-8')
    except Exception as e:
        return f"Failed to extract XML: {str(e)}"


def insert_header_near_text(doc_path: str, target_text: str, header_title: str, position: str = 'after', header_style: str = 'Heading 1') -> str:
    """Insert a header (with specified style) before or after the first paragraph containing target_text."""
    import os
    from docx import Document
    if not os.path.exists(doc_path):
        return f"Document {doc_path} does not exist"
    try:
        doc = Document(doc_path)
        found = False
        for i, para in enumerate(doc.paragraphs):
            if target_text in para.text:
                found = True
                # Create the new header paragraph with the specified style
                new_para = doc.add_paragraph(header_title, style=header_style)
                # Move the new paragraph to the correct position
                if position == 'before':
                    para._element.addprevious(new_para._element)
                else:
                    para._element.addnext(new_para._element)
                break
        if not found:
            return f"Target text '{target_text}' not found in document."
        doc.save(doc_path)
        return f"Header '{header_title}' (style: {header_style}) inserted {position} paragraph containing '{target_text}'."
    except Exception as e:
        return f"Failed to insert header: {str(e)}"


def insert_line_or_paragraph_near_text(doc_path: str, target_text: str, line_text: str, position: str = 'after', line_style: str = None) -> str:
    """
    Insert a new line or paragraph (with specified or matched style) before or after the first paragraph containing target_text.
    In Word, a new line is a new paragraph.
    """
    import os
    from docx import Document
    if not os.path.exists(doc_path):
        return f"Document {doc_path} does not exist"
    try:
        doc = Document(doc_path)
        found = False
        for i, para in enumerate(doc.paragraphs):
            if target_text in para.text:
                found = True
                # Determine style: use provided or match target
                style = line_style if line_style else para.style
                new_para = doc.add_paragraph(line_text, style=style)
                if position == 'before':
                    para._element.addprevious(new_para._element)
                else:
                    para._element.addnext(new_para._element)
                break
        if not found:
            return f"Target text '{target_text}' not found in document."
        doc.save(doc_path)
        return f"Line/paragraph inserted {position} paragraph containing '{target_text}' with style '{style}'."
    except Exception as e:
        return f"Failed to insert line/paragraph: {str(e)}"
